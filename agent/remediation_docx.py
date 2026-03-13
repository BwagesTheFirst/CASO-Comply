"""
CASO Comply -- Word Document (.docx) Accessibility Remediation Engine

Analyzes and remediates Word document accessibility issues using python-docx
and lxml for direct OOXML manipulation. Supports Gemini AI verification
for heading correction and alt text generation.
"""
from __future__ import annotations

import asyncio
import copy
import logging
import shutil
from collections import defaultdict
from pathlib import Path

import docx
from docx.shared import Pt
from lxml import etree

logger = logging.getLogger(__name__)

# OOXML namespaces
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
}


# ---------------------------------------------------------------------------
# 1. Analysis
# ---------------------------------------------------------------------------

def analyze_docx(file_path: str) -> dict:
    """Analyze Word document for accessibility issues. Returns score dict."""
    doc = docx.Document(file_path)

    # Check headings
    headings = []
    paragraphs_with_large_font = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1]) if para.style.name != "Heading" else 1
            headings.append({"level": level, "text": para.text})
        elif para.text.strip():
            # Check for visual-only formatting (bold + large = probable heading)
            for run in para.runs:
                if run.bold and run.font.size and run.font.size >= Pt(14):
                    paragraphs_with_large_font.append(para.text[:80])
                    break

    has_headings = len(headings) > 0
    hierarchy_valid = _check_heading_hierarchy(headings)

    # Check images for alt text
    images = _find_images(doc)
    images_with_alt = sum(1 for img in images if img.get("alt_text"))
    total_images = len(images)
    alt_text_coverage = images_with_alt / max(total_images, 1)

    # Check tables for headers
    tables = []
    for table in doc.tables:
        has_header = _table_has_header(table)
        tables.append({"rows": len(table.rows), "cols": len(table.columns), "has_header": has_header})
    tables_with_headers = sum(1 for t in tables if t["has_header"])
    total_tables = len(tables)
    table_header_coverage = tables_with_headers / max(total_tables, 1)

    # Check document language
    has_lang = _has_document_language(doc)

    # Check document title
    has_title = bool(doc.core_properties.title)

    # Check for floating images
    floating_images = _count_floating_images(doc)
    has_no_floating = floating_images == 0

    # Compute score
    score = _compute_docx_score(
        has_headings=has_headings,
        hierarchy_valid=hierarchy_valid,
        alt_text_coverage=alt_text_coverage,
        table_header_coverage=table_header_coverage,
        has_lang=has_lang,
        has_title=has_title,
        has_no_floating=has_no_floating,
        total_images=total_images,
        total_tables=total_tables,
    )

    return {
        "score": score,
        "headings": headings,
        "heading_issues": paragraphs_with_large_font,
        "images": {"total": total_images, "with_alt": images_with_alt},
        "tables": {"total": total_tables, "with_headers": tables_with_headers},
        "has_lang": has_lang,
        "has_title": has_title,
        "floating_images": floating_images,
        "page_count": _estimate_page_count(doc),
    }


def _check_heading_hierarchy(headings: list[dict]) -> bool:
    """Check that heading levels don't skip (e.g., H1 -> H3 with no H2)."""
    if not headings:
        return True
    prev_level = 0
    for h in headings:
        if h["level"] > prev_level + 1 and prev_level > 0:
            return False
        prev_level = h["level"]
    return True


def _find_images(doc) -> list[dict]:
    """Find all images in the document and check for alt text."""
    images = []

    # Check alt text via XML
    body = doc.element.body
    for drawing in body.iter("{%s}drawing" % NSMAP["w"]):
        docPr = drawing.find(".//{%s}docPr" % NSMAP["wp"])
        if docPr is not None:
            alt = docPr.get("descr", "")
            title = docPr.get("title", "")
            images.append({"alt_text": alt or title, "name": docPr.get("name", "")})

    return images


def _table_has_header(table) -> bool:
    """Check if a table has the header row repeat flag set."""
    tbl = table._tbl
    # Check first row for trPr/tblHeader
    first_row = tbl.find("{%s}tr" % NSMAP["w"])
    if first_row is None:
        return False
    trPr = first_row.find("{%s}trPr" % NSMAP["w"])
    if trPr is None:
        return False
    tblHeader = trPr.find("{%s}tblHeader" % NSMAP["w"])
    return tblHeader is not None


def _has_document_language(doc) -> bool:
    """Check if document language is set (in body text or style defaults)."""
    # Check body text
    body = doc.element.body
    for lang in body.iter("{%s}lang" % NSMAP["w"]):
        if lang.get("{%s}val" % NSMAP["w"]):
            return True
    # Check style defaults (where _set_document_language writes it)
    styles_element = doc.styles.element
    docDefaults = styles_element.find(f"{{{NSMAP['w']}}}docDefaults")
    if docDefaults is not None:
        for lang in docDefaults.iter("{%s}lang" % NSMAP["w"]):
            if lang.get("{%s}val" % NSMAP["w"]):
                return True
    return False


def _count_floating_images(doc) -> int:
    """Count images with text wrapping (non-inline)."""
    count = 0
    body = doc.element.body
    for drawing in body.iter("{%s}drawing" % NSMAP["w"]):
        anchor = drawing.find("{%s}anchor" % NSMAP["wp"])
        if anchor is not None:
            count += 1
    return count


def _estimate_page_count(doc) -> int:
    """Rough page estimate based on paragraph count."""
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    return max(1, para_count // 30 + 1)


def _compute_docx_score(
    has_headings: bool,
    hierarchy_valid: bool,
    alt_text_coverage: float,
    table_header_coverage: float,
    has_lang: bool,
    has_title: bool,
    has_no_floating: bool,
    total_images: int,
    total_tables: int,
) -> dict:
    """Compute 0-100 accessibility score for Word document."""
    checks = {
        "headings": {
            "passed": has_headings and hierarchy_valid,
            "weight": 20,
            "description": "Headings present with valid hierarchy",
        },
        "alt_text": {
            "passed": alt_text_coverage >= 0.9 if total_images > 0 else True,
            "weight": 25,
            "description": "Images have alt text" if total_images > 0 else "No images -- alt text not needed",
        },
        "table_headers": {
            "passed": table_header_coverage >= 0.9 if total_tables > 0 else True,
            "weight": 15,
            "description": "Tables have header rows" if total_tables > 0 else "No tables -- headers not needed",
        },
        "language": {
            "passed": has_lang,
            "weight": 10,
            "description": "Document language specified",
        },
        "title": {
            "passed": has_title,
            "weight": 10,
            "description": "Document title set",
        },
        "no_floating": {
            "passed": has_no_floating,
            "weight": 10,
            "description": "No floating images (inline only)",
        },
        "contrast": {
            "passed": True,  # Placeholder -- contrast check is complex
            "weight": 10,
            "description": "Color contrast adequate (not yet checked)",
        },
    }

    total = sum(c["weight"] for c in checks.values())
    earned = sum(c["weight"] for c in checks.values() if c["passed"])
    score = round(100 * earned / total) if total > 0 else 0

    return {"score": score, "checks": checks}


# ---------------------------------------------------------------------------
# 2. Remediation
# ---------------------------------------------------------------------------

def _classify_paragraphs(doc) -> list[dict]:
    """Classify paragraphs into heading/body tags by font size heuristic."""
    MAX_HEADING_CHARS = 200
    blocks = []

    font_sizes = set()
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        # Get effective font size
        size = None
        for run in para.runs:
            if run.font.size:
                size = run.font.size
                break
        if size is None:
            if para.style and para.style.font and para.style.font.size:
                size = para.style.font.size

        # Already a heading style
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1]) if para.style.name != "Heading" else 1
            blocks.append({
                "index": i,
                "text": para.text[:120],
                "type": f"H{level}",
                "font_size": size.pt if size else 0,
                "is_heading_style": True,
            })
        else:
            blocks.append({
                "index": i,
                "text": para.text[:120],
                "type": "P",
                "font_size": size.pt if size else 0,
                "is_heading_style": False,
            })
            if size:
                font_sizes.add(size.pt)

    if not font_sizes:
        return blocks

    # Identify probable headings: bold + large font + short text
    sorted_sizes = sorted(font_sizes, reverse=True)
    for block in blocks:
        if block["is_heading_style"]:
            continue
        if len(block["text"]) > MAX_HEADING_CHARS:
            continue
        if block["font_size"] <= 0:
            continue

        size_rank = sorted_sizes.index(block["font_size"]) if block["font_size"] in sorted_sizes else -1
        if size_rank == 0 and block["font_size"] >= 18:
            block["type"] = "H1"
        elif size_rank == 1 and block["font_size"] >= 14:
            block["type"] = "H2"
        elif size_rank == 2 and block["font_size"] >= 12:
            block["type"] = "H3"

    return blocks


def _apply_heading_styles(doc, classifications: list[dict]):
    """Apply proper Word heading styles based on classifications."""
    style_map = {"H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"}
    paragraphs = list(doc.paragraphs)

    for cls in classifications:
        if cls["type"] in style_map and not cls["is_heading_style"]:
            idx = cls["index"]
            if idx < len(paragraphs):
                paragraphs[idx].style = doc.styles[style_map[cls["type"]]]


def _set_table_headers(doc):
    """Set the tblHeader flag on the first row of every table."""
    w_ns = NSMAP["w"]
    for table in doc.tables:
        first_row = table._tbl.find(f"{{{w_ns}}}tr")
        if first_row is None:
            continue
        trPr = first_row.find(f"{{{w_ns}}}trPr")
        if trPr is None:
            trPr = etree.SubElement(first_row, f"{{{w_ns}}}trPr")
            first_row.insert(0, trPr)
        existing = trPr.find(f"{{{w_ns}}}tblHeader")
        if existing is None:
            etree.SubElement(trPr, f"{{{w_ns}}}tblHeader")


def _set_document_language(doc, lang: str = "en-US"):
    """Set the document-level language."""
    styles_element = doc.styles.element
    docDefaults = styles_element.find(f"{{{NSMAP['w']}}}docDefaults")
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_element, f"{{{NSMAP['w']}}}docDefaults")
        styles_element.insert(0, docDefaults)

    rPrDefault = docDefaults.find(f"{{{NSMAP['w']}}}rPrDefault")
    if rPrDefault is None:
        rPrDefault = etree.SubElement(docDefaults, f"{{{NSMAP['w']}}}rPrDefault")

    rPr = rPrDefault.find(f"{{{NSMAP['w']}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(rPrDefault, f"{{{NSMAP['w']}}}rPr")

    lang_elem = rPr.find(f"{{{NSMAP['w']}}}lang")
    if lang_elem is None:
        lang_elem = etree.SubElement(rPr, f"{{{NSMAP['w']}}}lang")

    lang_elem.set(f"{{{NSMAP['w']}}}val", lang)
    lang_elem.set(f"{{{NSMAP['w']}}}eastAsia", lang)
    lang_elem.set(f"{{{NSMAP['w']}}}bidi", "ar-SA")


def _set_alt_text_on_images(doc, alt_texts: dict):
    """Write AI-generated alt text to images.

    alt_texts: dict mapping image index to alt text string.
    """
    body = doc.element.body
    drawings = list(body.iter("{%s}drawing" % NSMAP["w"]))

    for idx, drawing in enumerate(drawings):
        alt = alt_texts.get(idx)
        if not alt:
            continue
        docPr = drawing.find(".//{%s}docPr" % NSMAP["wp"])
        if docPr is not None:
            docPr.set("descr", alt)


# ---------------------------------------------------------------------------
# 3. Main Pipeline
# ---------------------------------------------------------------------------

async def remediate_docx_async(
    file_path: str,
    output_path: str | None = None,
    verify: bool = True,
    *,
    gemini_provider: str = "standard",
    gcp_project: str = "",
    gcp_location: str = "us-central1",
) -> dict:
    """Full Word document remediation pipeline."""
    file_path = str(file_path)
    if output_path is None:
        p = Path(file_path)
        output_path = str(p.parent / f"{p.stem}_remediated{p.suffix}")

    logger.info("Analyzing input DOCX: %s", file_path)
    before_analysis = analyze_docx(file_path)

    # Open document for remediation
    doc = docx.Document(file_path)

    # Heuristic classification
    classifications = _classify_paragraphs(doc)

    # Apply heuristic fixes
    _apply_heading_styles(doc, classifications)
    _set_table_headers(doc)

    if not doc.core_properties.title:
        for cls in classifications:
            if cls["type"] == "H1":
                doc.core_properties.title = cls["text"][:256]
                break

    if not _has_document_language(doc):
        _set_document_language(doc)

    # Gemini AI verification
    verification_info = None
    if verify:
        from gemini_verify import verify_and_correct_office

        logger.info("Running Gemini AI verification on DOCX")
        tag_assignments = [
            {"mcid": i, "page": 0, "type": cls["type"], "text": cls["text"],
             "bbox": [0, 0, 0, 0], "font_size": cls["font_size"]}
            for i, cls in enumerate(classifications)
        ]

        verification_result = await verify_and_correct_office(
            file_path=file_path,
            structure_data=tag_assignments,
            format="docx",
            gemini_provider=gemini_provider,
            gcp_project=gcp_project,
            gcp_location=gcp_location,
        )

        verification_info = {
            "issues_found": verification_result.get("issues_found", []),
            "verification_score": verification_result.get("verification_score", 0.0),
        }

        # Apply Gemini corrections to headings
        corrected = verification_result.get("corrected_tags", [])
        if corrected:
            style_map = {"H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"}
            paragraphs = list(doc.paragraphs)
            non_empty_idx = 0
            for para in paragraphs:
                if not para.text.strip():
                    continue
                for ct in corrected:
                    if ct.get("mcid") == non_empty_idx:
                        tag_type = ct.get("type", "P")
                        if tag_type in style_map:
                            para.style = doc.styles[style_map[tag_type]]
                        elif tag_type == "P" and para.style.name.startswith("Heading"):
                            para.style = doc.styles["Normal"]
                        break
                non_empty_idx += 1

        # Apply AI alt texts
        alt_texts_raw = verification_result.get("alt_texts", {})
        if isinstance(alt_texts_raw, dict):
            flat_alts = {}
            for page_alts in alt_texts_raw.values():
                for entry in page_alts:
                    flat_alts[entry.get("image_index", 0)] = entry.get("alt_text", "")
            _set_alt_text_on_images(doc, flat_alts)
        elif isinstance(alt_texts_raw, list):
            flat_alts = {i: entry.get("description", "") for i, entry in enumerate(alt_texts_raw)}
            _set_alt_text_on_images(doc, flat_alts)

    # Save remediated document
    doc.save(output_path)
    logger.info("Saved remediated DOCX: %s", output_path)

    # Analyze the output
    after_analysis = analyze_docx(output_path)

    result = {
        "before": before_analysis,
        "after": after_analysis,
        "output_path": output_path,
        "page_count": before_analysis.get("page_count", 1),
    }

    if verification_info:
        result["verification"] = verification_info

    return result
